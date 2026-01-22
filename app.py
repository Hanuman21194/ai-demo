
import streamlit as st
from test_function_calling_streaming import tool_calling
from docx import Document
from io import BytesIO

# load_dotenv()
# client = OpenAI()
st.set_page_config(
    page_title ="GenAI Assistant",
    page_icon= "🤖",
    layout ="centered"
)
st.title ("🤖 GenAI Assistant")
st.caption ("Powered by OpenAI Response API (GPT-4o)")

user_prompt =st.text_area(
    "Enter your question:",
    height =150,
    placeholder ="Example: Explian Apache Spark architechture")

def generate_word_doc (question:str,answer:str) -> BytesIO : 
      """
     Create a Word document in memory and return it as BytesIO
      """
      doc =Document()
      doc.add_heading("AI Assistence Response",level=1)  
      
      doc.add_heading("Question",level=2)
      doc.add_paragraph(question)

      doc.add_heading("Answer",level=2)
      doc.add_paragraph(answer)
      
      buffer =BytesIO()
      doc.save(buffer)
      buffer.seek(0)
      return buffer
if st.button ("Ask AI"):
            if not user_prompt.strip():
                  st.warning("⚠️ Please enter a prompt")
            else :
                  st.subheader ("AI Response")
                  response_container =st.empty()
                  #full_response =""
                  with st.spinner("Thinking ...."):
                        try:  
                              #confidence = None     
                      #old way ---- # for chunk in tool_calling(user_prompt):
                              #       full_response  +=  chunk
                              #       response_container.markdown(full_response + "|")
                                 
                              # response_container.markdown(full_response)       
                                     #st.write (full_response)
                            

                              #response =tool_calling(user_prompt)
                              #st.write (response)
                              full_response=st.write_stream(tool_calling(user_prompt))
                              if full_response:
                                     st.write("---")
                              word_file = generate_word_doc(user_prompt,full_response)
                              st.download_button (
                                     label="📄 Download as Word",
                                     data =word_file,
                                     file_name="gen_ai_response.docx",
                                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                              )
                        except Exception as e:
                              st.error(f" Error {e}")            
