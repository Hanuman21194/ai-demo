import streamlit as st
import requests
from docx import Document
from io import BytesIO

import os
from dotenv import load_dotenv

load_dotenv()

SYSTEM_PROMPT = os.getenv(
    "SYSTEM_PROMPT",
    "You are a helpful AI assistant."
)

st.set_page_config(page_title="AI Assistant", layout="wide")

# --- 1. SESSION STATE (The "Memory") ---
if "messages" not in st.session_state:
    st.session_state.messages = []  # Stores chat history



# --- 2. WORD DOC UTILITY ---
def generate_word_doc(messages):
    doc = Document()
    doc.add_heading("AI Chat Export", level=1)
    for msg in messages:
        doc.add_heading(msg["role"].capitalize(), level=2)
        doc.add_paragraph(msg["content"])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

MODEL_OPTIONS = [
    "gpt-4o",
    "gpt-4.1",
    "gpt-4o-mini",
    "gpt-4.1-mini"
]

# Initialize session state properly
if "selected_model" not in st.session_state:
    st.session_state.selected_model = MODEL_OPTIONS[0]

col1, col2 = st.columns([3, 1])

with col2:
    st.session_state.selected_model = st.selectbox(
        "Model",
        MODEL_OPTIONS,
        index=MODEL_OPTIONS.index(st.session_state.selected_model),
        label_visibility="collapsed"
    )



# --- 2. SIDEBAR: PROMPT TRACKER ---
with st.sidebar:
    st.header("📝 Prompt History")
    st.caption("All user prompts in this session")


    user_prompts = [
        msg["content"]
        for msg in st.session_state.messages
        if msg["role"] == "user"
    ]

    if user_prompts:
        for idx, prompt_text in enumerate(reversed(user_prompts), start=1):
            st.markdown(f"**{len(user_prompts) - idx + 1}.** {prompt_text}")
    else:
        st.info("No prompts yet")

    st.divider()

    if st.button("🗑️ Clear All"):
        st.session_state.messages.clear()
        st.rerun()

    if st.session_state.messages:
        st.download_button(
            "📄 Export Full Conversation",
            data=generate_word_doc(st.session_state.messages),
            file_name="chat_history.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- 3. DISPLAY CHAT HISTORY ---
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# --- 4. CHAT INPUT (Gemini-Style) ---
if prompt := st.chat_input("Ask me anything..."):
  ############system prompt addition############
    if not any(msg["role"] == "system" for msg in st.session_state.messages):
        st.session_state.messages.insert(0, {"role": "system", "content": SYSTEM_PROMPT})

    # Display user message
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # Display assistant response with streaming
    with st.chat_message("assistant"):
        response_placeholder = st.empty()
        full_response = ""
        
        r = None
        try:
            # Call your FastAPI backend
            url = "http://127.0.0.1:8000/ask"
            r = requests.post(
                url,
                json={"messages": st.session_state.messages, "model": st.session_state.selected_model,
                      "system_prompt": SYSTEM_PROMPT},
                stream=True,
                timeout=60,
            )
            r.raise_for_status()
            for chunk in r.iter_content(chunk_size=1024, decode_unicode=True):
                if chunk:
                    full_response += chunk
                    response_placeholder.markdown(full_response + "▌")

            response_placeholder.markdown(full_response)  # Remove cursor
            st.session_state.messages.append({"role": "assistant", "content": full_response})

        except Exception as e:
            st.error(f"Backend Connection Error: {e}")
        finally:
            if r is not None:
                try:
                    r.close()
                except Exception:
                    pass

# # --- 5. SIDEBAR TOOLS ---
# with st.sidebar:
#     st.title("Settings & Tools")
#     if st.button("Clear Chat"):
#         st.session_state.messages.append({"role": "user", "content": prompt})
#         st.session_state.messages = []
#         st.rerun()
    
    